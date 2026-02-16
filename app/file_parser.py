"""文件解析模块 - 从 docx/pdf 提取文本内容。"""

from pathlib import Path

import docx
import pdfplumber


def parse_docx(file_path: Path) -> str:
    """从 Word 文档提取文本。

    Args:
        file_path: docx 文件路径

    Returns:
        提取的文本内容
    """
    doc = docx.Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def parse_pdf(file_path: Path) -> str:
    """从 PDF 提取文本。

    Args:
        file_path: PDF 文件路径

    Returns:
        提取的文本内容
    """
    text_parts = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

            # 提取表格
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cells = [str(cell).strip() for cell in row if cell]
                    if cells:
                        text_parts.append(" | ".join(cells))

    return "\n\n".join(text_parts)


def parse_file(file_path: Path) -> str:
    """根据文件类型提取文本。

    Args:
        file_path: 文件路径

    Returns:
        提取的文本内容

    Raises:
        ValueError: 不支持的文件格式
    """
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return parse_docx(file_path)
    elif suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")
